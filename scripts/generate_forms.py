import os
import sys
import openpyxl
import subprocess
import shutil
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml

sys.stdout.reconfigure(encoding='utf-8')

# --- CODE 39 PATTERNS ---
CODE39_PATTERNS = {
    '0': '000110100', '1': '100100001', '2': '001100001', '3': '101100000',
    '4': '000110001', '5': '100110000', '6': '001110000', '7': '000100101',
    '8': '100100100', '9': '001100100', 'A': '100001001', 'B': '001001001',
    'C': '101001000', 'D': '000011001', 'E': '100011000', 'F': '001011000',
    'G': '000001101', 'H': '100001100', 'I': '001001100', 'J': '000011100',
    'K': '100000011', 'L': '001000011', 'M': '101000010', 'N': '000010011',
    'O': '100010010', 'P': '001010010', 'Q': '000000111', 'R': '100000110',
    'S': '001000110', 'T': '000010110', 'U': '110000001', 'V': '011000001',
    'W': '111000000', 'X': '010010001', 'Y': '110010000', 'Z': '011010000',
    '-': '010000101', '.': '110000100', ' ': '011000100', '$': '010101000',
    '/': '010100010', '+': '010001010', '%': '000101010', '*': '010010100'
}

def generate_code39_svg(text, height=22, narrow_width=0.9, wide_width=2.2, show_text=True):
    full_text = f"*{text.upper()}*"
    rects = []
    current_x = 8.0  # Adequate quiet zone
    for char in full_text:
        pattern = CODE39_PATTERNS.get(char, CODE39_PATTERNS[' '])
        for i, bit in enumerate(pattern):
            is_bar = (i % 2 == 0)
            width = wide_width if bit == '1' else narrow_width
            if is_bar:
                rects.append(f'<rect x="{current_x:.1f}" y="0" width="{width:.1f}" height="{height}" fill="#000" />')
            current_x += width
        current_x += narrow_width  # gap
    current_x += 8.0  # Right quiet zone
    total_w = current_x
    text_svg = f'<text x="{total_w/2:.1f}" y="{height + 9}" font-family="Consolas, monospace" font-size="7.5" font-weight="bold" text-anchor="middle" fill="#000">*{text.upper()}*</text>' if show_text else ''
    total_h = height + (11 if show_text else 2)
    svg = f'''<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 {total_w:.1f} {total_h}" width="{total_w:.1f}" height="{total_h}">
        <rect width="100%" height="100%" fill="#fff"/>
        {''.join(rects)}
        {text_svg}
    </svg>'''
    return svg

def get_name_style(name):
    clean_name = name.replace(" ", "").replace("　", "")
    length = len(clean_name)
    if length <= 4:
        return 'font-size: 17px; letter-spacing: 3px;'
    elif length <= 6:
        return 'font-size: 15px; letter-spacing: 1.5px;'
    elif length <= 8:
        return 'font-size: 13.5px; letter-spacing: 1px;'
    elif length <= 11:
        return 'font-size: 12px; letter-spacing: 0.5px;'
    elif length <= 16:
        return 'font-size: 10.5px; letter-spacing: 0px;'
    else:
        return 'font-size: 9.5px; letter-spacing: -0.5px;'

def get_word_font_size(name):
    clean_name = name.replace(" ", "").replace("　", "")
    length = len(clean_name)
    if length <= 4:
        return Pt(14)
    elif length <= 6:
        return Pt(12.5)
    elif length <= 8:
        return Pt(11)
    elif length <= 11:
        return Pt(9.5)
    elif length <= 16:
        return Pt(8.5)
    else:
        return Pt(7.5)

import csv

def load_students(file_path):
    path = Path(file_path)
    students = []
    if path.suffix.lower() == '.csv':
        with open(path, 'r', encoding='utf-8-sig') as f:
            reader = csv.reader(f)
            rows = list(reader)
            for r in rows[1:]:
                if not r or not r[0].strip():
                    continue
                students.append({
                    'nichinoken_id': r[0].strip(),
                    'name': r[1].strip() if len(r) > 1 else '',
                    'kana': r[2].strip() if len(r) > 2 else '',
                    'class_name': r[3].strip() if len(r) > 3 else ''
                })
    else:
        wb = openpyxl.load_workbook(file_path)
        ws = wb.active
        rows = list(ws.iter_rows(values_only=True))
        for r in rows[1:]:
            if not r or not r[0]:
                continue
            students.append({
                'nichinoken_id': str(r[0]).strip(),
                'name': str(r[1]).strip() if r[1] else '',
                'kana': str(r[2]).strip() if len(r) > 2 and r[2] else '',
                'class_name': str(r[3]).strip() if len(r) > 3 and r[3] else ''
            })
    return students

def build_single_form_html(student, is_sample_filled=False, sample_info=None):
    barcode_svg = generate_code39_svg(student['nichinoken_id'])
    name_style = get_name_style(student['name'])
    
    check_no_change = ""
    check_has_change = ""
    change_content = ""
    change_reason = ""
    guardian_name = ""
    submit_date_m = "　"
    submit_date_d = "　"
    
    if is_sample_filled and sample_info:
        if sample_info.get('has_change'):
            check_has_change = '<span class="check-mark">✔</span>'
            change_content = f'<div class="handwritten">{sample_info.get("content", "")}</div>'
            change_reason = f'<div class="handwritten">{sample_info.get("reason", "")}</div>'
            guardian_name = f'<span class="handwritten">{sample_info.get("guardian", "")}</span>'
        else:
            check_no_change = '<span class="check-mark">✔</span>'
        submit_date_m = " 7 "
        submit_date_d = "15"
        
    return f'''
    <div class="form-page">
        <!-- Top cut line -->
        <div class="cut-line-container">
            <div class="cut-line-left"></div>
            <span class="cut-label">キリトリ</span>
            <div class="cut-line-right"></div>
        </div>

        <div class="form-wrapper">
            <!-- Header section with Barcode, Title and Date -->
            <div class="form-header">
                <div class="header-left">
                    <div class="barcode-box">
                        {barcode_svg}
                    </div>
                </div>
                <div class="header-center">
                    <h1 class="form-title">2026年度　夏期講習受講確認票</h1>
                </div>
                <div class="header-right">
                    <div class="submit-date">提出日：　{submit_date_m}　月　{submit_date_d}　日</div>
                </div>
            </div>

            <!-- Main Table -->
            <table class="form-table">
                <!-- Row 1: Student info -->
                <tr class="row-student-info">
                    <td class="cell-id">
                        <div class="cell-label">日能研番号</div>
                        <div class="cell-value-id">{student['nichinoken_id']}</div>
                    </td>
                    <td class="cell-name">
                        <div class="name-container">
                            <span class="cell-label-name">氏名</span>
                            <span class="cell-value-name" style="{name_style}">{student['name']}</span>
                            <span class="cell-honorific">様</span>
                        </div>
                    </td>
                    <td class="cell-class">
                        <div class="cell-label-class">5月度クラス</div>
                        <div class="cell-value-class">{student['class_name']}</div>
                    </td>
                </tr>

                <!-- Row 2: Notice -->
                <tr class="row-notice">
                    <td colspan="3" class="cell-notice">
                        ※3年予科教室生の夏期講習は4科目での受講を原則とします。2科目受講を希望される方は以下の「変更内容」にその旨をご記入ください。
                    </td>
                </tr>

                <!-- Row 3: No change option -->
                <tr class="row-option-no-change">
                    <td colspan="3" class="cell-option">
                        <div class="option-line">
                            <div class="checkbox-box">{check_no_change}</div>
                            <div class="option-text">夏期講習の受講内容に変更が<span class="badge-highlight">「ない」</span>（所属クラスの期間で　　科目を受講）</div>
                        </div>
                    </td>
                </tr>

                <!-- Row 4: Has change option -->
                <tr class="row-option-change">
                    <td class="cell-option-change-label">
                        <div class="option-line">
                            <div class="checkbox-box">{check_has_change}</div>
                            <div class="option-text">夏期講習の受講内容に変更が<span class="badge-highlight">「ある」</span></div>
                        </div>
                    </td>
                    <td colspan="2" class="cell-option-change-desc">
                        変更内容（科目、期間、他校舎でのご受講の場合、希望校舎など）を下記にご記入ください。
                    </td>
                </tr>

                <!-- Row 5: Change application details -->
                <tr class="row-application-details">
                    <td class="cell-vertical-header">
                        <div class="vertical-text">変更申請</div>
                    </td>
                    <td colspan="2" class="cell-application-body">
                        <div class="app-section">
                            <div class="section-title">《変更内容》</div>
                            <div class="section-content content-change">
                                {change_content}
                            </div>
                        </div>
                        <div class="app-section">
                            <div class="section-title">《変更理由》</div>
                            <div class="section-content content-reason">
                                {change_reason}
                            </div>
                        </div>
                        <div class="app-guardian-signature">
                            保護者氏名　<span class="signature-line">{guardian_name}</span>　<span class="seal-mark">印</span>
                        </div>
                    </td>
                </tr>
            </table>
        </div>
    </div>
    '''

def build_full_html(students, is_sample=False):
    samples_data = [
        {'has_change': False},
        {'has_change': False},
        {
            'has_change': True,
            'content': '第2期（8/1〜8/4） 横浜校にて受講を希望します。',
            'reason': '保護者の帰省日程と重複するため。',
            'guardian': '日能研 保護者'
        },
        {'has_change': False},
        {
            'has_change': True,
            'content': '4科目のうち「算数・国語」の2科目のみ受講を希望します。',
            'reason': '他習い事の全国大会遠征日程と重複するため。',
            'guardian': '保護者 太郎'
        },
        {'has_change': False}
    ]
    
    pages_html = []
    for i, s in enumerate(students):
        sample_info = samples_data[i % len(samples_data)] if is_sample else None
        pages_html.append(build_single_form_html(s, is_sample, sample_info))
        
    all_pages = "\n".join(pages_html)
    
    return f'''<!DOCTYPE html>
<html lang="ja">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>2026年度 夏期講習受講確認票 (A5サイズ) - 印刷プレビュー</title>
    <style>
        @import url('https://fonts.googleapis.com/css2?family=BIZ+UDPGothic:wght@400;700&family=Yuji+Boku&display=swap');

        @page {{
            size: A5 landscape;
            margin: 0;
        }}

        * {{
            box-sizing: border-box;
            margin: 0;
            padding: 0;
        }}

        body {{
            background-color: #3f444e;
            font-family: "BIZ UDPGothic", "Hiragino Kaku Gothic ProN", "Meiryo", sans-serif;
            color: #000;
            line-height: 1.25;
            -webkit-print-color-adjust: exact;
            print-color-adjust: exact;
        }}

        /* Toolbar */
        .preview-toolbar {{
            position: sticky;
            top: 0;
            left: 0;
            right: 0;
            background: #1e293b;
            color: #fff;
            padding: 10px 20px;
            display: flex;
            align-items: center;
            justify-content: space-between;
            box-shadow: 0 4px 12px rgba(0,0,0,0.3);
            z-index: 1000;
        }}

        .toolbar-title {{
            font-size: 14px;
            font-weight: 700;
            display: flex;
            align-items: center;
            gap: 8px;
        }}

        .toolbar-buttons {{
            display: flex;
            align-items: center;
            gap: 10px;
        }}

        .btn-toggle {{
            background: #334155;
            color: #f1f5f9;
            border: 1px solid #475569;
            padding: 6px 12px;
            font-size: 12px;
            border-radius: 6px;
            text-decoration: none;
            display: inline-flex;
            align-items: center;
            gap: 6px;
            transition: all 0.15s;
        }}
        .btn-toggle:hover {{
            background: #475569;
            color: #fff;
        }}

        .btn-print {{
            background: #2563eb;
            color: #fff;
            border: none;
            padding: 6px 16px;
            font-size: 12px;
            font-weight: 700;
            border-radius: 6px;
            cursor: pointer;
            display: inline-flex;
            align-items: center;
            gap: 6px;
            transition: all 0.15s;
        }}
        .btn-print:hover {{
            background: #1d4ed8;
            box-shadow: 0 2px 6px rgba(37,99,235,0.4);
        }}

        /* A5 Landscape Page Layout (210mm x 148mm) */
        .form-page {{
            width: 210mm;
            height: 148mm;
            margin: 20px auto;
            background: #ffffff;
            padding: 7mm 10mm 6mm 10mm;
            box-shadow: 0 6px 20px rgba(0,0,0,0.35);
            page-break-after: always;
            box-sizing: border-box;
            position: relative;
            display: flex;
            flex-direction: column;
            justify-content: flex-start;
        }}

        /* Cut line */
        .cut-line-container {{
            display: flex;
            align-items: center;
            width: 100%;
            margin-bottom: 6px;
        }}

        .cut-line-left, .cut-line-right {{
            flex-grow: 1;
            border-top: 1px dashed #666;
        }}

        .cut-label {{
            padding: 0 10px;
            font-size: 9.5px;
            color: #333;
            letter-spacing: 3px;
            font-weight: 500;
        }}

        /* Form Wrapper */
        .form-wrapper {{
            width: 100%;
            flex-grow: 1;
            display: flex;
            flex-direction: column;
        }}

        /* Header Grid: Barcode Left, Title Center, Date Right */
        .form-header {{
            display: grid;
            grid-template-columns: 110px 1fr 110px;
            align-items: flex-end;
            margin-bottom: 5px;
            padding: 0 2px;
        }}

        .header-left {{
            text-align: left;
        }}

        .barcode-box {{
            display: inline-block;
        }}

        .header-center {{
            text-align: center;
        }}

        .form-title {{
            font-size: 16.5px;
            font-weight: 900;
            letter-spacing: 2px;
            color: #000;
            white-space: nowrap;
        }}

        .header-right {{
            text-align: right;
            padding-bottom: 2px;
        }}

        .submit-date {{
            font-size: 10px;
            font-weight: bold;
            letter-spacing: 1px;
            white-space: nowrap;
        }}

        /* Main Table */
        .form-table {{
            width: 100%;
            border-collapse: collapse;
            border: 1.8px solid #000;
            flex-grow: 1;
        }}

        .form-table td {{
            border: 1px solid #000;
            vertical-align: middle;
        }}

        /* Row 1: Student info */
        .row-student-info {{
            height: 38px;
        }}

        .cell-id {{
            width: 22%;
            text-align: center;
            padding: 2px 4px;
            background: #fafafa;
        }}

        .cell-label {{
            font-size: 8px;
            color: #444;
            margin-bottom: 1px;
        }}

        .cell-value-id {{
            font-family: "Consolas", "Courier New", monospace;
            font-size: 13.5px;
            font-weight: bold;
            letter-spacing: 0.5px;
        }}

        .cell-name {{
            width: 55%;
            padding: 2px 8px;
            overflow: hidden;
        }}

        .name-container {{
            display: flex;
            align-items: baseline;
            justify-content: space-between;
            width: 100%;
            overflow: hidden;
        }}

        .cell-label-name {{
            font-size: 9px;
            color: #444;
            width: 24px;
            flex-shrink: 0;
        }}

        .cell-value-name {{
            font-weight: bold;
            text-align: center;
            flex-grow: 1;
            white-space: nowrap;
            overflow: hidden;
            text-overflow: clip;
        }}

        .cell-honorific {{
            font-size: 12px;
            font-weight: bold;
            margin-left: 4px;
            flex-shrink: 0;
        }}

        .cell-class {{
            width: 23%;
            text-align: center;
            padding: 2px 4px;
            background: #fafafa;
            white-space: nowrap;
        }}

        .cell-label-class {{
            font-size: 8px;
            color: #444;
            margin-bottom: 1px;
            white-space: nowrap;
        }}

        .cell-value-class {{
            font-size: 13.5px;
            font-weight: bold;
            white-space: nowrap;
        }}

        /* Row 2: Notice */
        .row-notice {{
            background-color: #f1f3f5;
            height: 20px;
        }}

        .cell-notice {{
            padding: 2px 6px;
            font-size: 7.8px;
            color: #222;
            line-height: 1.25;
        }}

        /* Row 3: Option No Change */
        .row-option-no-change {{
            height: 32px;
        }}

        .cell-option {{
            padding: 3px 10px;
        }}

        .option-line {{
            display: flex;
            align-items: center;
            gap: 8px;
        }}

        .checkbox-box {{
            display: flex;
            align-items: center;
            justify-content: center;
            width: 16px;
            height: 16px;
            border: 1.5px solid #000;
            background: #fff;
            flex-shrink: 0;
        }}

        .check-mark {{
            font-size: 15px;
            font-weight: 900;
            color: #000;
            line-height: 1;
            transform: translate(0.5px, -1.5px);
        }}

        .option-text {{
            font-size: 10.5px;
            font-weight: 500;
            white-space: nowrap;
        }}

        .badge-highlight {{
            background: #e2e8f0;
            border: 1px solid #94a3b8;
            padding: 0.5px 5px;
            border-radius: 2px;
            font-weight: bold;
            margin: 0 3px;
        }}

        /* Row 4: Option Has Change */
        .row-option-change {{
            height: 32px;
        }}

        .cell-option-change-label {{
            width: 40%;
            padding: 3px 10px;
            background: #fafafa;
        }}

        .cell-option-change-desc {{
            padding: 3px 8px;
            font-size: 8.5px;
            color: #333;
            line-height: 1.25;
        }}

        /* Row 5: Application Details */
        .row-application-details {{
            height: 145px;
        }}

        .cell-vertical-header {{
            text-align: center;
            background: #fafafa;
            width: 32px;
            padding: 8px 0;
            vertical-align: middle;
        }}

        .vertical-text {{
            writing-mode: vertical-rl;
            text-orientation: upright;
            font-size: 11.5px;
            font-weight: bold;
            letter-spacing: 6px;
            margin: 0 auto;
        }}

        .cell-application-body {{
            padding: 8px 12px;
            position: relative;
            vertical-align: top !important;
            height: 145px;
        }}

        .app-section {{
            margin-bottom: 8px;
        }}

        .section-title {{
            font-size: 9.5px;
            font-weight: bold;
            color: #000;
            margin-bottom: 2px;
        }}

        .section-content {{
            min-height: 24px;
            padding: 1px 4px;
            font-size: 11px;
        }}

        .handwritten {{
            font-family: "Yuji Boku", "Meiryo", cursive, sans-serif;
            font-size: 11.5px;
            color: #0f172a;
            letter-spacing: 0.5px;
        }}

        .app-guardian-signature {{
            position: absolute;
            bottom: 8px;
            right: 14px;
            font-size: 10.5px;
            font-weight: bold;
            display: flex;
            align-items: flex-end;
            gap: 4px;
        }}

        .signature-line {{
            display: inline-block;
            width: 170px;
            border-bottom: 1px solid #000;
            min-height: 16px;
            text-align: center;
            font-size: 11.5px;
        }}

        .seal-mark {{
            display: inline-flex;
            align-items: center;
            justify-content: center;
            width: 20px;
            height: 20px;
            border: 1px solid #777;
            border-radius: 50%;
            font-size: 9px;
            color: #555;
            margin-left: 4px;
        }}

        /* Print Media Queries */
        @media print {{
            body {{
                background: #fff;
            }}
            .preview-toolbar {{
                display: none !important;
            }}
            .form-page {{
                margin: 0;
                padding: 6mm 8mm;
                box-shadow: none;
                width: 100vw;
                height: 100vh;
                page-break-after: always;
            }}
        }}
    </style>
</head>
<body>

    <div class="preview-toolbar">
        <div class="toolbar-title">
            <span>📄</span>
            <span>2026年度 夏期講習受講確認票【A5サイズ】（{'✏️ 記入済テストサンプル' if is_sample else '📋 白紙受講票'}・全{len(students)}名分）</span>
        </div>
        <div class="toolbar-buttons">
            <a href="{'受講確認票_白紙.html' if is_sample else '受講確認票_記入済サンプル.html'}" class="btn-toggle">
                {'🔄 白紙版に切り替え' if is_sample else '✏️ 記入済テスト版に切り替え'}
            </a>
            <button onclick="window.print()" class="btn-print">
                <span>🖨️</span> A5印刷 / PDF保存
            </button>
        </div>
    </div>

    {all_pages}

</body>
</html>
'''

def safe_copy(src, dst):
    try:
        shutil.copy(src, dst)
        return True
    except Exception as e:
        print(f"Warning: Could not copy {src.name} to {dst}: {e}")
        return False

def generate_word_document(students, output_path, is_template=False):
    doc = docx.Document()
    
    # Page setup - A5 Landscape (210mm width x 148mm height)
    sections = doc.sections
    for section in sections:
        section.page_width = Mm(210)
        section.page_height = Mm(148)
        section.top_margin = Mm(8)
        section.bottom_margin = Mm(8)
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)

    target_students = students if not is_template else [{
        'nichinoken_id': '«日能研番号»',
        'name': '«氏名»',
        'kana': '«氏名カナ»',
        'class_name': '«クラス»'
    }]

    for idx, student in enumerate(target_students):
        if idx > 0:
            doc.add_page_break()

        # 1. Cut line
        p_cut = doc.add_paragraph()
        p_cut.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cut.paragraph_format.space_after = Pt(4)
        p_cut.paragraph_format.space_before = Pt(0)
        r_cut = p_cut.add_run("-------------------- キリトリ --------------------")
        r_cut.font.size = Pt(8)
        r_cut.font.color.rgb = RGBColor(120, 120, 120)

        # 2. Header Table (Barcode, Title, Date)
        header_table = doc.add_table(rows=1, cols=3)
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_table.autofit = False
        
        header_table.columns[0].width = Mm(35)
        header_table.columns[1].width = Mm(120)
        header_table.columns[2].width = Mm(35)

        # Cell 0: Barcode text
        c0 = header_table.cell(0, 0)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r0 = p0.add_run(f"*{student['nichinoken_id']}*")
        r0.font.name = "Consolas"
        r0.font.size = Pt(10)
        r0.bold = True

        # Cell 1: Title
        c1 = header_table.cell(0, 1)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run("2026年度 夏期講習受講確認票")
        r1.font.size = Pt(13)
        r1.bold = True

        # Cell 2: Date
        c2 = header_table.cell(0, 2)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = p2.add_run("提出日：　月　日")
        r2.font.size = Pt(8.5)

        # Spacer
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(1)
        p_space.paragraph_format.space_after = Pt(2)

        # 3. Main Form Table
        table = doc.add_table(rows=5, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        table.columns[0].width = Mm(40)
        table.columns[1].width = Mm(110)
        table.columns[2].width = Mm(40)

        # Row 0: Student Info
        # Left: ID
        r0_c0 = table.cell(0, 0)
        p = r0_c0.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("日能研番号\n")
        r.font.size = Pt(7.5)
        r = p.add_run(student['nichinoken_id'])
        r.font.size = Pt(11)
        r.bold = True

        # Center: Name
        r0_c1 = table.cell(0, 1)
        p = r0_c1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("氏名　")
        r.font.size = Pt(8.5)
        
        name_font_size = get_word_font_size(student['name']) if not is_template else Pt(11)
        r = p.add_run(student['name'])
        r.font.size = name_font_size
        r.bold = True
        
        r = p.add_run("　様")
        r.font.size = Pt(9.5)

        # Right: Class
        r0_c2 = table.cell(0, 2)
        p = r0_c2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("5月度クラス\n")
        r.font.size = Pt(7.5)
        r = p.add_run(student['class_name'])
        r.font.size = Pt(11)
        r.bold = True

        # Row 1: Notice (merged)
        r1_c0 = table.cell(1, 0)
        r1_c2 = table.cell(1, 2)
        r1_merged = r1_c0.merge(r1_c2)
        p = r1_merged.paragraphs[0]
        r = p.add_run("※3年予科教室生の夏期講習は4科目での受講を原則とします。2科目受講を希望される方は以下の「変更内容」にその旨をご記入ください。")
        r.font.size = Pt(7)

        # Row 2: No change (merged)
        r2_c0 = table.cell(2, 0)
        r2_c2 = table.cell(2, 2)
        r2_merged = r2_c0.merge(r2_c2)
        p = r2_merged.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("□　夏期講習の受講内容に変更が「ない」（所属クラスの期間で　　科目を受講）")
        r.font.size = Pt(9)

        # Row 3: Has change
        r3_c0 = table.cell(3, 0)
        p = r3_c0.paragraphs[0]
        r = p.add_run("□　夏期講習の受講内容に変更が「ある」")
        r.font.size = Pt(8.5)

        r3_c1 = table.cell(3, 1)
        r3_c2 = table.cell(3, 2)
        r3_merged = r3_c1.merge(r3_c2)
        p = r3_merged.paragraphs[0]
        r = p.add_run("変更内容（科目、期間、他校舎でのご受講の場合、希望校舎など）を下記にご記入ください。")
        r.font.size = Pt(7.5)

        # Row 4: Application details
        r4_c0 = table.cell(4, 0)
        p = r4_c0.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("\n変\n更\n申\n請\n")
        r.font.size = Pt(9.5)
        r.bold = True

        r4_c1 = table.cell(4, 1)
        r4_c2 = table.cell(4, 2)
        r4_merged = r4_c1.merge(r4_c2)
        p = r4_merged.paragraphs[0]
        r = p.add_run("《変更内容》\n\n《変更理由》\n\n")
        r.font.size = Pt(8)
        r.bold = True
        
        p_sig = r4_merged.add_paragraph()
        p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p_sig.add_run("保護者氏名 ＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿ 印")
        r.font.size = Pt(8.5)

        # Apply cell borders
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._element.get_or_add_tcPr()
                tcBorders = parse_xml(r'''
                    <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                        <w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>
                        <w:left w:val="single" w:sz="6" w:space="0" w:color="000000"/>
                        <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>
                        <w:right w:val="single" w:sz="6" w:space="0" w:color="000000"/>
                    </w:tcBorders>
                ''')
                tcPr.append(tcBorders)

    doc.save(output_path)
    print(f"Saved Word document (A5): {output_path}")

def main():
    root_dir = Path(__file__).resolve().parent.parent
    default_input = root_dir / "templates" / "students_template.csv"
    
    input_path = sys.argv[1] if len(sys.argv) > 1 else str(default_input)
    out_dir = root_dir / "output"
    out_dir.mkdir(parents=True, exist_ok=True)
    
    print(f"Loading students from: {input_path}")
    students = load_students(input_path)
    print(f"Loaded {len(students)} students")

    # 1. Generate HTML Blank Forms (A5 Landscape)
    html_blank = build_full_html(students, is_sample=False)
    html_blank_path = out_dir / "受講確認票_白紙.html"
    with open(html_blank_path, "w", encoding="utf-8") as f:
        f.write(html_blank)

    # 2. Generate HTML Sample Filled Forms (A5 Landscape)
    html_sample = build_full_html(students, is_sample=True)
    html_sample_path = out_dir / "受講確認票_記入済サンプル.html"
    with open(html_sample_path, "w", encoding="utf-8") as f:
        f.write(html_sample)

    # Also copy HTML to workspace root for easy access
    with open(root_dir / "受講確認票_印刷プレビュー.html", "w", encoding="utf-8") as f:
        f.write(html_blank)

    # 3. Generate Word Document (.docx)
    docx_all_path = out_dir / "受講確認票_全生徒分.docx"
    generate_word_document(students, docx_all_path, is_template=False)
    
    docx_template_path = out_dir / "受講確認票_差し込み印刷テンプレート.docx"
    generate_word_document(students, docx_template_path, is_template=True)

    # 4. Generate PDFs using Chrome headless (A5 Landscape)
    chrome_path = r"C:\Program Files\Google\Chrome\Application\chrome.exe"
    if not os.path.exists(chrome_path):
        chrome_path = r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"

    pdf_blank_path = out_dir / "受講確認票_全生徒一括.pdf"
    pdf_sample_path = out_dir / "受講確認票_テストスキャン用_記入済.pdf"

    if os.path.exists(chrome_path):
        print("Rendering A5 PDF via Chrome headless...")
        cmd_blank = [
            chrome_path,
            "--headless",
            "--disable-gpu",
            "--no-pdf-header-footer",
            f"--print-to-pdf={pdf_blank_path}",
            str(html_blank_path)
        ]
        subprocess.run(cmd_blank, check=True)

        cmd_sample = [
            chrome_path,
            "--headless",
            "--disable-gpu",
            "--no-pdf-header-footer",
            f"--print-to-pdf={pdf_sample_path}",
            str(html_sample_path)
        ]
        subprocess.run(cmd_sample, check=True)

    print("\n--- ALL A5 FORMS GENERATED SUCCESSFULLY ---")

if __name__ == "__main__":
    main()
